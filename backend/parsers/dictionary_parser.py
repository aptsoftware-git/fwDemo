"""
Dictionary parser for FRS Data Management System.
Parses Dictionary.docx to extract column mappings for different sheet types.

NOTE: Dictionary.docx parsing is optional. The system uses metadata_definitions.py
as the primary source of column metadata. Dictionary.docx is only parsed if available
and valid, but its absence will not cause failures.
"""
import os
import logging
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document

# Import structured metadata as primary source
try:
    from metadata_definitions import get_column_names, SHEET_METADATA
    METADATA_AVAILABLE = True
except ImportError:
    METADATA_AVAILABLE = False
    logging.warning("metadata_definitions.py not found - using fallback mode")

# Cache for dictionary data to avoid repeated file reads
_dictionary_cache: Optional[Dict] = None

logger = logging.getLogger(__name__)


def parse_dictionary(dictionary_path: str = None) -> Dict[str, List[str]]:
    """
    Get column mappings for each sheet type.
    
    This function prioritizes structured metadata from metadata_definitions.py.
    Dictionary.docx parsing is attempted only as a supplement, not a requirement.
    
    Args:
        dictionary_path: Path to Dictionary.docx file (optional). If None, skips file parsing.
    
    Returns:
        Dictionary mapping sheet types to list of column names in sequence.
        Example:
        {
            'APPX_A_BVEH': ['Column1', 'Column2', ...],
            'APPX_A_CVEH': ['Column1', 'Column2', ...],
            ...
        }
    
    Note:
        This function never raises exceptions. If Dictionary.docx cannot be parsed,
        it falls back to metadata_definitions.py.
    """
    global _dictionary_cache
    
    # Return cached data if available
    if _dictionary_cache is not None:
        return _dictionary_cache
    
    # Try to use structured metadata as primary source
    if METADATA_AVAILABLE:
        logger.info("Using structured metadata from metadata_definitions.py")
        column_mappings = {
            sheet_type: get_column_names(sheet_type)
            for sheet_type in ['APPX_A_BVEH', 'APPX_A_CVEH', 'ARMT', 'SA', 'INST', 'CBRN']
        }
        _dictionary_cache = column_mappings
        return column_mappings
    
    # If no structured metadata available, try Dictionary.docx as fallback
    logger.warning("metadata_definitions.py not available, attempting Dictionary.docx parsing")
    
    # Determine dictionary path
    if dictionary_path is None:
        # Default path: ../docs/Dictionary.docx relative to this file
        backend_dir = Path(__file__).resolve().parent.parent
        dictionary_path = backend_dir.parent / 'docs' / 'Dictionary.docx'
    
    dictionary_path = Path(dictionary_path)
    
    if not dictionary_path.exists():
        logger.warning(f"Dictionary.docx not found at: {dictionary_path}. Using default column mappings.")
        # Return default column mappings
        return _get_default_column_mappings()
    
    try:
        doc = Document(str(dictionary_path))
        
        # Dictionary to store column mappings
        column_mappings = {
            'APPX_A_BVEH': [],
            'APPX_A_CVEH': [],
            'ARMT': [],
            'SA': [],
            'INST': [],
            'CBRN': []
        }
        
        current_sheet_type = None
        
        # Parse document paragraphs and tables
        for element in doc.paragraphs:
            text = element.text.strip()
            
            # Identify sheet type headers (case-insensitive matching)
            if 'appx a' in text.lower() and 'b veh' in text.lower():
                current_sheet_type = 'APPX_A_BVEH'
            elif 'appx a' in text.lower() and 'c veh' in text.lower():
                current_sheet_type = 'APPX_A_CVEH'
            elif 'armt' in text.lower() and 'appx' not in text.lower():
                current_sheet_type = 'ARMT'
            elif text.upper() == 'SA' or 'small arms' in text.lower():
                current_sheet_type = 'SA'
            elif 'inst' in text.lower():
                current_sheet_type = 'INST'
            elif 'cbrn' in text.lower():
                current_sheet_type = 'CBRN'
        
        # Parse tables for column names
        for table in doc.tables:
            # Attempt to extract column names from table headers
            if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
                header_row = table.rows[0]
                columns = [cell.text.strip() for cell in header_row.cells if cell.text.strip()]
                
                # Try to identify which sheet type this table belongs to
                # This is a heuristic - you may need to adjust based on actual document structure
                if current_sheet_type and columns:
                    column_mappings[current_sheet_type] = columns
        
        # If no tables found or mapping is incomplete, use default mappings
        for sheet_type, columns in column_mappings.items():
            if not columns:
                logger.warning(f"No columns found for {sheet_type} in Dictionary.docx, using defaults")
                column_mappings[sheet_type] = _get_default_columns_for_sheet(sheet_type)
        
        # Cache the result
        _dictionary_cache = column_mappings
        
        return column_mappings
    
    except Exception as e:
        logger.error(f"Failed to parse Dictionary.docx: {str(e)}. Using default column mappings.")
        return _get_default_column_mappings()


def _get_default_columns_for_sheet(sheet_type: str) -> List[str]:
    """
    Get default column names for a sheet type.
    
    Args:
        sheet_type: Sheet type identifier
    
    Returns:
        List of default column names
    """
    # Try to use metadata_definitions if available
    if METADATA_AVAILABLE:
        return get_column_names(sheet_type)
    
    # Fallback to generic column names
    return [f'Column_{i+1}' for i in range(20)]


def _get_default_column_mappings() -> Dict[str, List[str]]:
    """
    Get default column mappings for all sheet types.
    
    Returns:
        Dictionary of sheet Type to column names
    """
    sheet_types = ['APPX_A_BVEH', 'APPX_A_CVEH', 'ARMT', 'SA', 'INST', 'CBRN']
    return {
        sheet_type: _get_default_columns_for_sheet(sheet_type)
        for sheet_type in sheet_types
    }


def get_column_mapping(sheet_type: str, dictionary_path: str = None) -> List[str]:
    """
    Get column mapping for a specific sheet type.
    
    Args:
        sheet_type: Sheet type identifier (e.g., 'APPX_A_BVEH')
        dictionary_path: Optional path to Dictionary.docx
    
    Returns:
        List of column names in sequence
    
    Raises:
        ValueError: If sheet_type is invalid
    """
    mappings = parse_dictionary(dictionary_path)
    
    if sheet_type not in mappings:
        raise ValueError(f"Invalid sheet type: {sheet_type}. Valid types: {list(mappings.keys())}")
    
    return mappings[sheet_type]


def clear_cache():
    """Clear the dictionary cache. Useful for testing or if dictionary file is updated."""
    global _dictionary_cache
    _dictionary_cache = None


# For testing and debugging
if __name__ == "__main__":
    try:
        mappings = parse_dictionary()
        print("Dictionary parsed successfully!")
        print("\nColumn mappings:")
        for sheet_type, columns in mappings.items():
            print(f"\n{sheet_type}:")
            print(f"  {len(columns)} columns: {columns[:5]}..." if len(columns) > 5 else f"  {columns}")
    except Exception as e:
        print(f"Error: {e}")
